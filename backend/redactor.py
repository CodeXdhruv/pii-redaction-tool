import re
import docx
from docx import Document
from presidio_analyzer import AnalyzerEngine, PatternRecognizer, Pattern
from presidio_analyzer.nlp_engine import NlpEngineProvider
from faker import Faker

class PIIRedactor:
    def __init__(self, spacy_model="en_core_web_sm"):
        # Configure spaCy engine for Presidio
        configuration = {
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "en", "model_name": spacy_model}]
        }
        provider = NlpEngineProvider(nlp_configuration=configuration)
        nlp_engine = provider.create_engine()
        
        # Optimize spaCy pipeline by disabling unused components (tagger, parser, etc.)
        # This speeds up spaCy processing significantly (up to 3x-4x faster)
        if "en" in nlp_engine.nlp:
            nlp = nlp_engine.nlp["en"]
            for pipe in ["parser", "tagger", "lemmatizer", "attribute_ruler"]:
                if nlp.has_pipe(pipe):
                    nlp.remove_pipe(pipe)
        
        # Initialize Presidio Analyzer
        self.analyzer = AnalyzerEngine(nlp_engine=nlp_engine)
        
        # Add custom recognizers
        self._add_custom_recognizers()
        
        # Initialize Faker for replacements
        self.faker = Faker()
        Faker.seed(42)  # For deterministic replacements across runs if desired
        
        # Cache to map original PII value -> fake replacement to preserve consistency
        self.pii_cache = {}
        
        # Cache for PII analysis results of identical text blocks (e.g. repeated headers/footers)
        self.analysis_cache = {}
        
        # Simple entity name mapping for display/reporting
        self.entity_mapping = {
            "PERSON": "Full Name",
            "EMAIL_ADDRESS": "Email Address",
            "PHONE_NUMBER": "Phone Number",
            "ORGANIZATION": "Company Name",
            "LOCATION": "Physical Address",
            "US_SSN": "Social Security Number",
            "CREDIT_CARD": "Credit Card Number",
            "DATE_TIME": "Date of Birth / Date",
            "IP_ADDRESS": "IP Address"
        }

    def _add_custom_recognizers(self):
        # 1. Custom Phone Number Recognizer (matching Indian +91 and 10-digit formats)
        indian_phone_pattern = Pattern(
            name="indian_phone_pattern",
            regex=r"(?:\+91[\s-]?)?[6789]\d{9}\b|\b\d{5}[-\s]\d{5}\b",
            score=0.95
        )
        phone_recognizer = PatternRecognizer(
            supported_entity="PHONE_NUMBER",
            patterns=[indian_phone_pattern],
            context=["phone", "mobile", "contact", "tel", "telephone", "cell"]
        )
        self.analyzer.registry.add_recognizer(phone_recognizer)

        # 2. SSN Pattern Recognizer (adds a backup regex for standard SSNs)
        ssn_pattern = Pattern(
            name="ssn_pattern",
            regex=r"\b\d{3}-\d{2}-\d{4}\b",
            score=0.95
        )
        ssn_recognizer = PatternRecognizer(
            supported_entity="US_SSN",
            patterns=[ssn_pattern],
            context=["ssn", "social security", "sec number"]
        )
        self.analyzer.registry.add_recognizer(ssn_recognizer)

        # 3. Custom Company Name Recognizer (capitalized terms ending in Limited, Bank, Securities, Ltd, etc.)
        company_pattern = Pattern(
            name="company_pattern",
            regex=r"\b[A-Z][a-zA-Z0-9]*(?:\s+[A-Z][a-zA-Z0-9]*)*\s+(?:Limited|Ltd\.?|Bank|Securities|Corporation|Pvt\.?\s+Ltd\.?|Co\.)\b",
            score=0.95
        )
        company_recognizer = PatternRecognizer(
            supported_entity="ORGANIZATION",
            patterns=[company_pattern],
            context=["company", "firm", "bank", "manager", "auditor", "underwriter"]
        )
        self.analyzer.registry.add_recognizer(company_recognizer)

        # 4. Custom Date / Date of Birth Recognizer (matching textual and numeric date formats)
        date_pattern1 = Pattern(
            name="date_pattern_text",
            regex=r"\b\d{1,2}[-/\s](?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-zA-Z]*[-/\s]\d{2,4}\b",
            score=0.95
        )
        date_pattern2 = Pattern(
            name="date_pattern_num",
            regex=r"\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b",
            score=0.95
        )
        date_recognizer = PatternRecognizer(
            supported_entity="DATE_TIME",
            patterns=[date_pattern1, date_pattern2],
            context=["birth", "dob", "born", "date", "dated"]
        )
        self.analyzer.registry.add_recognizer(date_recognizer)

        # 5. Custom Credit Card Recognizer (ensuring 100% recall with higher score to avoid spaCy PERSON confusion)
        cc_pattern = Pattern(
            name="cc_pattern",
            regex=r"\b\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}\b",
            score=0.95
        )
        cc_recognizer = PatternRecognizer(
            supported_entity="CREDIT_CARD",
            patterns=[cc_pattern],
            context=["card", "credit card", "visa", "mastercard"]
        )
        self.analyzer.registry.add_recognizer(cc_recognizer)

    def get_fake_replacement(self, text: str, entity_type: str) -> str:
        """Returns a consistent fake replacement for a given PII text."""
        # Check cache first
        cache_key = (text.strip().lower(), entity_type)
        if cache_key in self.pii_cache:
            return self.pii_cache[cache_key]

        # Generate fake alternative
        fake_val = ""
        if entity_type == "PERSON":
            fake_val = self.faker.name()
        elif entity_type == "EMAIL_ADDRESS":
            fake_val = self.faker.email()
        elif entity_type == "PHONE_NUMBER":
            # Generate phone matching original length/style if possible, or standard +91
            fake_val = f"+91 {self.faker.numerify('##########')}"
        elif entity_type == "ORGANIZATION":
            fake_val = self.faker.company()
        elif entity_type == "LOCATION":
            fake_val = self.faker.address().replace("\n", ", ")
        elif entity_type == "US_SSN":
            fake_val = self.faker.ssn()
        elif entity_type == "CREDIT_CARD":
            fake_val = self.faker.credit_card_number()
        elif entity_type == "DATE_TIME":
            # Check if it looks like a year, or full date
            if len(text) == 4 and text.isdigit():
                fake_val = str(self.faker.random_int(1970, 2005))
            else:
                fake_val = self.faker.date_of_birth(minimum_age=18, maximum_age=90).strftime("%d-%B-%Y")
        elif entity_type == "IP_ADDRESS":
            fake_val = self.faker.ipv4()
        else:
            fake_val = f"<{entity_type}_REDACTED>"

        self.pii_cache[cache_key] = fake_val
        return fake_val

    def analyze_text(self, text: str):
        """Analyzes a text block, filters by confidence, and resolves overlapping entities."""
        cleaned_text = text.strip()
        if not cleaned_text or len(cleaned_text) < 4:
            return []
        
        # Check cache
        if cleaned_text in self.analysis_cache:
            return self.analysis_cache[cleaned_text]
        
        # Analyze with Presidio
        results = self.analyzer.analyze(
            text=text,
            language="en",
            entities=list(self.entity_mapping.keys())
        )
        
        # Filter results with confidence threshold (e.g. >= 0.4)
        filtered = [r for r in results if r.score >= 0.4]
        
        # Resolve conflicts/overlaps
        # Sort by score (descending), then by span length (descending)
        filtered_sorted = sorted(
            filtered,
            key=lambda x: (x.score, x.end - x.start),
            reverse=True
        )
        
        non_overlapping = []
        for res in filtered_sorted:
            # Check if this overlaps with any already accepted result
            overlap = False
            for accepted in non_overlapping:
                if max(res.start, accepted.start) < min(res.end, accepted.end):
                    overlap = True
                    break
            if not overlap:
                non_overlapping.append(res)
                
        # Sort by start position (ascending) for clean processing
        final_results = sorted(non_overlapping, key=lambda x: x.start)
        
        # Cache results
        self.analysis_cache[cleaned_text] = final_results
        return final_results

    def redact_text_block(self, text: str) -> tuple[str, list[dict]]:
        """Redacts a single string, returning the redacted string and redaction metadata."""
        results = self.analyze_text(text)
        if not results:
            return text, []

        # Sort matches in descending order to perform safe replacements
        results_desc = sorted(results, key=lambda x: x.start, reverse=True)
        
        redacted_text = text
        changes = []
        
        for r in results_desc:
            original_val = text[r.start:r.end]
            fake_val = self.get_fake_replacement(original_val, r.entity_type)
            
            # Perform replacement in the string
            redacted_text = redacted_text[:r.start] + fake_val + redacted_text[r.end:]
            
            changes.append({
                "original": original_val,
                "replacement": fake_val,
                "entity_type": self.entity_mapping.get(r.entity_type, r.entity_type),
                "confidence": r.score,
                "start": r.start,
                "end": r.end
            })
            
        return redacted_text, changes[::-1] # return changes in original ascending order

    def redact_paragraph(self, paragraph, stats):
        """Redacts a docx paragraph at the run level to preserve original formatting."""
        text = paragraph.text
        if not text.strip():
            return

        results = self.analyze_text(text)
        if not results:
            return

        # Sort results in descending order of start position
        results_desc = sorted(results, key=lambda x: x.start, reverse=True)
        runs = paragraph.runs
        
        # Map paragraph characters to run index and character offset inside that run
        char_to_run = []
        for run_idx, run in enumerate(runs):
            for char_idx in range(len(run.text)):
                char_to_run.append((run_idx, char_idx))

        # Check if the text length matches char_to_run list length
        if len(char_to_run) != len(text):
            # Fallback if there is a mismatch (e.g. due to fields/tabs XML difference)
            # Reconstruct the paragraph with a single run or replace paragraph text directly
            # This is safer than throwing an error or mismatching offsets
            for r in results_desc:
                original_val = text[r.start:r.end]
                fake_val = self.get_fake_replacement(original_val, r.entity_type)
                text = text[:r.start] + fake_val + text[r.end:]
                
                stats.append({
                    "original": original_val,
                    "replacement": fake_val,
                    "entity_type": self.entity_mapping.get(r.entity_type, r.entity_type),
                    "confidence": r.score
                })
            paragraph.text = text
            return

        for r in results_desc:
            start = r.start
            end = r.end
            original_val = text[start:end]
            fake_val = self.get_fake_replacement(original_val, r.entity_type)

            r_start_idx, char_start_offset = char_to_run[start]
            r_end_idx, char_end_offset = char_to_run[end - 1]

            # Clear out runs between r_start_idx and r_end_idx
            for idx in range(r_start_idx + 1, r_end_idx):
                runs[idx].text = ""

            if r_start_idx == r_end_idx:
                # Update text in the single run
                run_text = runs[r_start_idx].text
                new_run_text = run_text[:char_start_offset] + fake_val + run_text[char_end_offset + 1:]
                runs[r_start_idx].text = new_run_text
            else:
                # Split replacement between start and end runs
                run_start_text = runs[r_start_idx].text
                run_end_text = runs[r_end_idx].text
                
                runs[r_start_idx].text = run_start_text[:char_start_offset] + fake_val
                runs[r_end_idx].text = run_end_text[char_end_offset + 1:]

            # We do not need to rebuild char_to_run because we are processing matches in descending 
            # order (right-to-left). Any replacement at index 'start' does not affect the run index 
            # and offset mappings for characters located *before* 'start'.
            pass

            stats.append({
                "original": original_val,
                "replacement": fake_val,
                "entity_type": self.entity_mapping.get(r.entity_type, r.entity_type),
                "confidence": r.score
            })

    def redact_document(self, input_path: str, output_path: str) -> list[dict]:
        """Redacts an entire docx file (paragraphs, tables, headers, footers)."""
        doc = Document(input_path)
        stats = []

        # 1. Redact headers/footers
        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    self.redact_paragraph(p, stats)
                for table in section.header.tables:
                    self.redact_table(table, stats)
            if section.footer:
                for p in section.footer.paragraphs:
                    self.redact_paragraph(p, stats)
                for table in section.footer.tables:
                    self.redact_table(table, stats)

        # 2. Redact main body paragraphs
        for p in doc.paragraphs:
            self.redact_paragraph(p, stats)

        # 3. Redact tables
        for table in doc.tables:
            self.redact_table(table, stats)

        doc.save(output_path)
        return stats

    def redact_table(self, table, stats):
        """Recursively redact table cells."""
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    self.redact_paragraph(p, stats)
                for nested_table in cell.tables:
                    self.redact_table(nested_table, stats)
